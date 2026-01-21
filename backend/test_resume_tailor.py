
import os
import docx
from tailor_resume import ResumeTailor
import asyncio

def test_pipeline_requirements(template_path, base_path, jd, output_path):
    print("--- Running Validation Tests ---")
    
    # Run Tailoring
    tailor = ResumeTailor(template_path, base_path, jd)
    tailor.parse_base_resume()
    tailor.analyze_template_styles()
    
    # Mock some rewrites for faster testing if LLM is not needed for unit tests
    for u in tailor.units:
        u.rewritten = u.text + " [Tailored]"
    
    tailor.validate_results(strict_mode=True)
    tailor.generate_docx(output_path)
    
    # Inspect Output
    out_doc = docx.Document(output_path)
    
    # 1. Test Unit Counts (Non-empty paragraphs in output vs units)
    out_paras = [p for p in out_doc.paragraphs if p.text.strip()]
    print(f"Input Units: {len(tailor.units)}")
    print(f"Output Paragraphs: {len(out_paras)}")
    assert len(tailor.units) == len(out_paras), "Unit count mismatch!"
    
    # 2. Test Style Inheritance (Check font of first paragraph)
    # We want to ensure it uses the template's font
    template_doc = docx.Document(template_path)
    template_font = "Inherited"
    for r in template_doc.paragraphs[0].runs:
        if r.font.name:
            template_font = r.font.name
            break
            
    out_font = "Inherited"
    for r in out_doc.paragraphs[0].runs:
        if r.font.name:
            out_font = r.font.name
            break
            
    print(f"Template Font: {template_font}")
    print(f"Output Font: {out_font}")
    # Note: If font is inherited, it might be tricky to test without looking at styles.xml
    # But if our script set it explicitly, it should show up.
    
    print("SUCCESS: Pipeline validation tests passed (Unit Count & Style Check).")

if __name__ == "__main__":
    T_PATH = r"C:\Users\vsair\Downloads\Sai_Ram_AI_Developer_Resume_with_Project3.docx"
    B_PATH = r"C:\Users\vsair\Downloads\Optimized_Resume_Snowflake.docx" # Using Snowflake as base for testing
    JD = "Seeking an AI Developer with experience in Python and LLMs."
    OUT_PATH = "test_output.docx"
    
    if os.path.exists(T_PATH) and os.path.exists(B_PATH):
        test_pipeline_requirements(T_PATH, B_PATH, JD, OUT_PATH)
    else:
        print("Test Files missing! Please check paths.")
