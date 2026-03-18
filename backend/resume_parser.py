"""
Resume Parser - Extracts text from PDF and DOCX files
"""

import io
import logging
from typing import Optional, Union, Dict, Any

logger = logging.getLogger(__name__)


async def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file content
    
    Args:
        file_content: Raw bytes of the PDF file
        
    Returns:
        Extracted text from the PDF
    """
    try:
        import PyPDF2
        
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = "\n".join(text_parts)
        
        if not full_text.strip():
            logger.warning("PDF extraction returned empty text - might be image-based PDF")
            return ""
            
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> dict:
    """
    Extract text and metadata from DOCX file content
    
    Returns:
        Dictionary with 'text' and 'metadata' (font_family, etc.)
    """
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        detected_font = None
        detected_size = None
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                
                # Try to detect font and size from the first few meaningful paragraphs
                if not detected_font or not detected_size:
                    # Check runs
                    for run in paragraph.runs:
                        if not detected_font and run.font.name:
                            detected_font = run.font.name
                        if not detected_size and run.font.size:
                            detected_size = int(run.font.size.pt)
                        if detected_font and detected_size:
                            break
                    # Check style if runs didn't have it
                    if not detected_font and paragraph.style and paragraph.style.font and paragraph.style.font.name:
                        detected_font = paragraph.style.font.name
                    if not detected_size and paragraph.style and paragraph.style.font and paragraph.style.font.size:
                        detected_size = int(paragraph.style.font.size.pt)

        # Final fallback - check document defaults if still not found
        if not detected_font:
            try:
                # Check Normal style font
                if 'Normal' in doc.styles and doc.styles['Normal'].font.name:
                    detected_font = doc.styles['Normal'].font.name
            except Exception:
                pass
        
        if not detected_size:
            try:
                if 'Normal' in doc.styles and doc.styles['Normal'].font.size:
                    detected_size = int(doc.styles['Normal'].font.size.pt)
            except Exception:
                pass
        
        # If still nothing, defaults
        if not detected_font:
            detected_font = "Arial"
        if not detected_size:
            detected_size = 11

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return {
            "text": "\n".join(text_parts),
            "metadata": {
                "font_family": detected_font,
                "font_size": detected_size
            }
        }
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


async def parse_resume(file_content: bytes, filename: str) -> Union[str, dict]:
    """
    Parse resume file and extract text and metadata based on file type
    
    Args:
        file_content: Raw bytes of the file
        filename: Original filename to determine type
        
    Returns:
        Either the extracted text (str) or a dict with 'text' and 'metadata'
    """
    from typing import Union
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        # PDF currently only returns text
        text = await extract_text_from_pdf(file_content)
        return {"text": text, "metadata": {}}
    elif filename_lower.endswith('.docx'):
        return await extract_text_from_docx(file_content)
    elif filename_lower.endswith('.doc'):
        raise ValueError("Old .doc format not supported. Please convert to .docx or .pdf")
    elif filename_lower.endswith('.txt'):
        text = file_content.decode('utf-8', errors='ignore')
        return {"text": text, "metadata": {}}
    else:
        raise ValueError(f"Unsupported file type. Please upload PDF, DOCX, or TXT file")



def validate_magic_numbers(file_content: bytes, filename: str) -> bool:
    """
    Validate file content against magic numbers for its extension.
    
    Args:
        file_content: Raw bytes of the file
        filename: The filename to check extension against
        
    Returns:
        True if valid, False otherwise
    """
    filename_lower = filename.lower()
    
    # PDF Magic Number: %PDF- (25 50 44 46 2D)
    if filename_lower.endswith('.pdf'):
        return file_content.startswith(b'%PDF-')
        
    # DOCX Magic Number: PK (50 4B 03 04)
    elif filename_lower.endswith('.docx'):
        return file_content.startswith(b'PK\x03\x04')
        
    # TXT has no magic number, but we can check for binary content
    elif filename_lower.endswith('.txt'):
        try:
            file_content.decode('utf-8')
            return True
        except UnicodeDecodeError:
            return False
            
    return False


def validate_resume_file(filename: str, file_content: bytes) -> Optional[str]:
    """
    Validate resume file before processing
    
    Args:
        filename: The filename
        file_content: Raw bytes of the file
        
    Returns:
        Error message if invalid, None if valid
    """
    # Check file extension
    valid_extensions = ['.pdf', '.docx', '.txt']
    has_valid_ext = any(filename.lower().endswith(ext) for ext in valid_extensions)
    
    if not has_valid_ext:
        return "Invalid file type. Please upload PDF, DOCX, or TXT file"
    
    file_size = len(file_content)
    
    # Check file size (max 5MB - reduced from 10MB for security)
    max_size = 5 * 1024 * 1024  # 5MB
    if file_size > max_size:
        return "File too large. Maximum size is 5MB"
    
    # Check file size (min 10 bytes)
    min_size = 10
    if file_size < min_size:
        return f"Resume file looks empty ({file_size} bytes). Minimum required is {min_size} bytes."
        
    # Check magic numbers
    if not validate_magic_numbers(file_content, filename):
        return "Invalid file content. The file does not match its extension."
    
    return None


