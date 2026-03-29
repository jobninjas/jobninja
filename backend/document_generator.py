"""
Document Generator - Creates optimized resumes and cover letters
"""

import os
import re
import io
import logging
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aiohttp
import asyncio
import json
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, Union

from resume_parser_deterministic import (
    parse_resume_deterministic,
    merge_with_ai_facts,
    validate_parse,
)

from resume_analyzer import call_groq_api, unified_api_call, clean_json_response, extract_resume_data

logger = logging.getLogger(__name__)

# Add file handler for persistent AI debug logs
try:
    log_path = os.path.join(os.path.dirname(__file__), "ai_debug.log")
    fh = logging.FileHandler(log_path)
    fh.setLevel(logging.INFO)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    fh.setFormatter(formatter)
    logger.addHandler(fh)
except Exception as e:
    print(f"Failed to set up file logging in document_generator: {e}")

# These are now imported from resume_analyzer.py
# GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
# GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
# GROQ_MODEL = "llama-3.3-70b-versatile"

COVER_LETTER_AND_COLD_EMAIL_PROMPT = """
You are a JSON API. Output only a valid JSON object. No prose. Start with {{

╔══════════════════════════════════════════════════════════════╗
║  CRITICAL ROLE CLARIFICATION — READ BEFORE ANYTHING ELSE   ║
╠══════════════════════════════════════════════════════════════╣
║  SENDER:   {name} — the JOB SEEKER writing these messages  ║
║  RECEIVER: The hiring manager / recruiter at {target_company} ║
║                                                              ║
║  {name} IS WRITING TO APPLY FOR A JOB.                     ║
║  DO NOT write as if a recruiter is contacting {name}.       ║
║  DO NOT address {name} in the message.                      ║
║  DO NOT say "Hi {name}" or "Dear {name}".                   ║
║  The message is FROM {name}, TO the hiring team.            ║
╚══════════════════════════════════════════════════════════════╝

SENDER (writing these messages):
  Name:     {name}
  Email:    {email}
  Phone:    {phone}
  Role:     {current_title}

RECIPIENT (receiving these messages):
  Company:  {target_company}
  Role applying for: {target_role}
  Hiring Manager: {hiring_manager} (use "Hiring Team" if unknown)

════════════════════════════════════════════
CANDIDATE FACTS (LOCKED — never invent or change):
════════════════════════════════════════════
Name: {name}
Email: {email}
Phone: {phone}
Location: {location}
Current Title: {current_title}
Years of Experience: {years_exp}

Top 3 technical achievements from resume (use these as evidence):
{achievement_1}
{achievement_2}
{achievement_3}

Projects (real names only):
{project_list}

════════════════════════════════════════════
TARGET JOB:
═════════════════════════════
Role: {target_role}
Company: {target_company}
Location: {target_location}
Mission: {target_mission} (What they do)
Keywords: {target_keywords}
Pain Points: {target_pain_points} (e.g., scaling, user retention, data quality)

════════════════════════════════════════════
VARIATIONS STRATEGY:
════════════════════════════════════════════
Variation A: HIGH ALIGNMENT (The "I am the solution" approach)
Focus: Mapping the top 3 achievements directly to the target role's pain points.
Tone: Confident, expert, solution-oriented.

Variation B: COMPANY VISION (The "I am a believer" approach)
Focus: Connecting the candidate’s project history to the company’s specific mission and long-term goals.
Tone: Enthusiastic, visionary, culture-fit.

════════════════════════════════════════════
RULES FOR COVER LETTER:
════════════════════════════════════════════
1. Professional, high-standard headers.
2. 3-4 dense, high-signal paragraphs.
3. NO brackets. NO [Your Name]. NO [Company Name]. Use the target values provided.
4. Format: Plain text, professional line breaks.

════════════════════════════════════════════
RULES FOR COLD EMAIL:
════════════════════════════════════════════
1. Subject line included.
2. Max 100 words.
3. Hook-Benefit-Ask structure.
4. NO brackets.

════════════════════════════════════════════
OUTPUT JSON SCHEMA:
════════════════════════════════════════════
{{
  "cover_letter_A": "full text...",
  "cover_letter_B": "full text...",
  "cold_email_A": "full text...",
  "cold_email_B": "full text..."
}}
"""

EXPERT_SYSTEM_PROMPT = """
You are an expert resume writer specializing in AI/ML and software engineering roles in the US job market. Your job is to rewrite and tailor a candidate's resume to match a specific job description WITHOUT removing, summarizing, or shortening any experience. 

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
‼️ RULE ZERO — BEFORE ANYTHING ELSE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
NEVER start your output with "Dear", "Hello", "Hi", or any greeting.
NEVER write a cover letter. NEVER write an introduction paragraph.
NEVER write "Dear Hiring Team,".
NEVER invent companies, job titles, dates, or people.

STRICT RULES — never break these:
1. NEVER shrink, summarize, or remove bullet points from any job. Every role must have at least 6-8 detailed bullet points.
2. NEVER reduce the total length of the resume. The output must be LONGER than the input, not shorter.
3. Every job must include a "Project Summary" paragraph describing the company and the candidate's role context.
4. Every job must end with an "Environment:" line listing all tools and technologies used in that role.
5. Bullet points must be BOLD and detailed — minimum 25 words each.
6. Add or strengthen keywords from the job description naturally inside existing bullets. Do NOT invent fake experience.
7. Rewrite the Professional Summary to directly mirror the job description language and requirements.
8. Skills section must be a full categorized list — minimum 8 categories.
9. "name" field MUST ONLY contain the candidate's full legal name as found in the source resume. NEVER append job titles or designations to the name.
10. Output must be structured as valid JSON.
11. Never output a shortened or "clean" version. The goal is a COMPREHENSIVE, ATS-optimized resume.
"""

EXPERT_TAILORING_USER_PROMPT = """
Tailor this resume for the job description below. Follow every rule in the system prompt strictly.

=== CANDIDATE RESUME ===
{raw_resume_text}

=== JOB DESCRIPTION ===
{job_description}

=== OUTPUT FORMAT (Return JSON only) ===
{{
  "tailored_resume": {{
    "name": "Full Name",
    "title": "Exact Job Title from JD",
    "contact": {{
       "email": "email",
       "phone": "phone",
       "location": "location"
    }},
    "summary": ["sentence 1", "sentence 2", "sentence 3", "sentence 4"],
    "skills": {{
       "AI/ML": [],
       "Languages": [],
       "Data Engineering": [],
       "Cloud & MLOps": [],
       "GenAI & LLMs": [],
       "Databases": [],
       "DevOps": [],
       "Frameworks": []
    }},
    "experience": [
      {{
        "company": "Company",
        "title": "Title",
        "location": "Location",
        "start": "Mon YYYY",
        "end": "Mon YYYY",
        "project_summary": "2-3 sentences describing the company and candidate's context.",
        "bullets": [
           "**Bold 25+ word detailed bullet 1...**",
           "**Bold 25+ word detailed bullet 2...**",
           "**Bold 25+ word detailed bullet 3...**",
           "**Bold 25+ word detailed bullet 4...**",
           "**Bold 25+ word detailed bullet 5...**",
           "**Bold 25+ word detailed bullet 6...**"
        ],
        "environment": "tool1, tool2, tool3..."
      }}
    ],
    "projects": [
      {{
        "name": "Name",
        "tech": "Stack",
        "bullets": ["detailed bullet 1", "detailed bullet 2"]
      }}
    ],
    "certifications": [],
    "education": [
      {{
        "degree": "Degree",
        "major": "Major",
        "university": "University",
        "year": "Year"
      }}
    ]
  }},
  "cover_letter": "Detailed cover letter text"
}}
"""

# ============================================
# STRUCTURED RESUME SCHEMA (Step 3 & 4)
# ============================================

class ResumeHeader(BaseModel):
    full_name: str
    city_state: str
    phone: str
    email: str
    linkedin: str = ""
    portfolio: str = ""

class CoreSkills(BaseModel):
    languages: List[str] = []
    data_etl: List[str] = []
    cloud: List[str] = []
    databases: List[str] = []
    devops_tools: List[str] = []
    other: List[str] = []

class ExperienceRole(BaseModel):
    company: str
    job_title: str
    city_state_or_remote: str
    start: str
    end: str
    project_summary: str = ""
    bullets: List[str]
    environment: str = ""

class ProjectItem(BaseModel):
    name: str
    tech_stack: List[str] = []
    link: str = ""
    bullets: List[str]

class EducationItem(BaseModel):
    degree: Optional[str] = ""
    major: Optional[str] = ""
    university: Optional[str] = ""
    year: Optional[str] = ""

class ResumeDataSchema(BaseModel):
    """The structured data for a tailored resume"""
    header: ResumeHeader
    target_title: str
    positioning_statement: str
    core_skills: CoreSkills
    experience: List[ExperienceRole]
    projects: List[ProjectItem] = []
    education: List[EducationItem] = []
    certifications: List[str] = []

class ExpertTailoringOutput(BaseModel):
    """The complete response from the Expert AI"""
    alignment_highlights: List[str]
    cover_letter: str
    resume_data: ResumeDataSchema


def cleanup_bullet(s: str) -> str:
    """Removes weird bullet characters and extra whitespace (Step 4)"""
    import re
    cleaned = re.sub(r'[•●▪︎\-]', '', str(s))
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


def extract_json_from_response(text: str) -> dict:
    """
    Robustly extract JSON from AI response that may contain surrounding text.
    Tries multiple strategies in order.
    """
    if not text or not isinstance(text, str):
        return {}
    
    import json
    import re
    text = text.strip()
    
    # Strategy 1: Response IS pure JSON
    if text.startswith('{'):
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            # Try to find where the valid JSON ends
            pass
    
    # Strategy 2: JSON inside ```json ... ``` code block
    code_block = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', text)
    if code_block:
        try:
            return json.loads(code_block.group(1))
        except:
            pass
    
    # Strategy 3: Find the largest {...} block in the response
    # Walk from first { to find matching }
    start = text.find('{')
    if start != -1:
        depth = 0
        for i, ch in enumerate(text[start:], start):
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    candidate = text[start:i+1]
                    try:
                        return json.loads(candidate)
                    except:
                        break
    
    # Strategy 4: Response is free text — return empty dict to trigger fallback
    logger.error(f"Could not extract JSON from response. First 300 chars: {text[:300]}")
    return {}


def strip_appended_skill_suffixes(text: str) -> str:
    """
    Programmatically remove the suffix-appending bug pattern.
    Catches all variants the AI produces regardless of prompt instructions.
    """
    # Patterns: bullet ends with ", utilizing/leveraging/applying X."
    suffix_patterns = [
        r',\s*utilizing\s+[A-Za-z0-9\s,/\-\.]+(?:for [^.]+)?\.',
        r',\s*leveraging\s+[A-Za-z0-9\s,/\-\.]+(?:for [^.]+)?\.',
        r',\s*applying\s+principles\s+of\s+[^.]+\.',
        r',\s*ensuring\s+seamless\s+integration\s+with\s+[^.]+\.',
        r',\s*with\s+a\s+focus\s+on\s+building\s+highly\s+scalable[^.]*\.',
        r',\s*with\s+expertise\s+in\s+[^.]+\.',
        r',\s*with\s+a\s+strong\s+[^.]+\.',
    ]

    if not text or not isinstance(text, str):
        return text or ""

    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if not line:
            cleaned_lines.append("")
            continue
        original = line
        for pattern in suffix_patterns:
            # Only match at end of line (after real content)
            match = re.search(pattern + r'\s*$', line, re.IGNORECASE)
            if match and match.start() > 20:  # guard: don't truncate very short lines
                line = line[:match.start()].rstrip().rstrip(',') + '.'
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)


def strip_summary_prefix(text: str) -> str:
    """
    Remove the 'Job Title — ' prefix from summary and fix literal \n characters.
    """
    if not text or not isinstance(text, str):
        return text or ""
        
    # Remove "AI Engineer — " or similar from summary start
    text = re.sub(r'^[^\n\u2014]+[\u2014]\s*', '', text)
    # Fix literal \n or \\n from AI
    return text.replace('\\n', '\n').replace('\\\\n', '\n').strip()
    


def extract_company_from_title(title: str) -> tuple:
    """
    Parses "AI Engineer at Astra Corp" → ("AI Engineer", "Astra Corp")
    or "AI Engineer | Astra Corp" → ("AI Engineer", "Astra Corp")
    Returns (clean_title, company) — company is '' if not parseable.
    """
    if not title:
        return ("", "")
    
    for separator in [" at ", " @ ", " | ", " - "]:
        if separator in title:
            parts = title.split(separator, 1)
            return (parts[0].strip(), parts[1].strip())
    return (title.strip(), "")


def prepare_tailoring_facts(resume_json: dict) -> dict:
    """
    Extracts and formats facts from resume_json for the expert tailoring prompt.
    Includes original highlights/bullets to enable reframing.
    """
    person = resume_json.get("person", {})
    name = person.get("fullName") or f"{person.get('firstName', '')} {person.get('lastName', '')}".strip() or "Candidate"
    
    facts = {
        "name": name,
        "email": person.get("email", "N/A"),
        "phone": person.get("phone", "N/A"),
        "location": person.get("location", "N/A"),
        "employer_count": len(resume_json.get("employment_history", [])),
        "employer_list": "",
        "project_list": "",
        "cert_list": "",
        "edu_list": ""
    }
    
    # Format Employer List (with raw bullets for AI reframing)
    employers = []
    for emp in resume_json.get("employment_history", []):
        if not isinstance(emp, dict): continue
        comp = emp.get("company", "Unknown")
        title = emp.get("title", "Unknown")
        loc = emp.get("location", "N/A")
        start = emp.get("startDate", "")
        end = emp.get("endDate", "Present")
        highlights = emp.get("highlights") or emp.get("description", [])
        if isinstance(highlights, str):
            highlights = [highlights]
        
        bullet_text = "\n".join([f"  * {h}" for h in highlights if h and isinstance(h, str)])
        employers.append(f"- {comp} | {title} | {loc} ({start} to {end})\n  ORIGINAL HIGHLIGHTS:\n{bullet_text}")
    facts["employer_list"] = "\n".join(employers)
    
    # Format Project List
    projects = []
    raw_projects = resume_json.get("projects", [])
    if not isinstance(raw_projects, list): raw_projects = []
    if not raw_projects and isinstance(resume_json.get("skills"), dict):
        raw_projects = resume_json.get("skills", {}).get("projects", [])
        if not isinstance(raw_projects, list): raw_projects = []
        
    for proj in raw_projects:
        if not isinstance(proj, dict): continue
        p_name = proj.get("name") or "Unnamed Project"
        tech = proj.get("tech_stack", [])
        tech_str = ", ".join(tech) if isinstance(tech, list) else str(tech)
        p_bullets = proj.get("highlights") or proj.get("bullets", [])
        if isinstance(p_bullets, str):
            p_bullets = [p_bullets]
        p_bullet_text = "\n".join([f"  * {b}" for b in p_bullets if b and isinstance(b, str)])
        projects.append(f"- {p_name} (Tech: {tech_str})\n  ORIGINAL HIGHLIGHTS:\n{p_bullet_text}")
    facts["project_list"] = "\n".join(projects)
    
    # Format Certifications
    certs = []
    skills = resume_json.get("skills", {})
    raw_certs = skills.get("certifications", []) if isinstance(skills, dict) else []
    if not isinstance(raw_certs, list): raw_certs = []
    if not raw_certs:
        raw_certs = resume_json.get("certifications", [])
    if isinstance(raw_certs, list):
        for cert in raw_certs:
            if cert: certs.append(f"- {cert}")
    facts["cert_list"] = "\n".join(certs)
    
    # Format Education
    edus = []
    for edu in resume_json.get("education", []):
        if not isinstance(edu, dict): continue
        degree = str(edu.get("degree", "")).strip()
        major = str(edu.get("major", "")).strip()
        school = str(edu.get("school", "") or edu.get("university", "")).strip()
        year = str(edu.get("graduationDate") or edu.get("year", "")).strip()
        
        # Clean "undefined" leaks in education
        fields = [degree, major, school, year]
        clean_fields = [f if str(f).lower() not in ["undefined", "null", "none"] else "" for f in fields]
        degree, major, school, year = clean_fields
        
        if degree or major or school:
            edus.append(f"- {degree}, {major} from {school} ({year})")
    facts["edu_list"] = "\n".join(edus)
    
    # Final name cleanup in facts
    facts["name"] = re.sub(r'(?i)undefined|null|none', '', str(facts.get("name", ""))).strip()
    if not facts["name"]:
        facts["name"] = "Candidate"
        
    return facts


def prepare_cover_letter_cold_email_facts(resume_json: dict, job_description: str, job_title: str, company: str) -> dict:
    """
    Extracts and formats facts for the COVER_LETTER_AND_COLD_EMAIL_PROMPT.
    """
    person = resume_json.get("person", {})
    name = person.get("fullName") or f"{person.get('firstName', '')} {person.get('lastName', '')}".strip() or "Candidate"
    
    # Calculate years of experience
    total_years = 0
    for emp in resume_json.get("employment_history", []):
        try:
            start = emp.get("startDate")
            end = emp.get("endDate") or "2026-03" # Current date approx
            if start and end:
                s_y = int(start.split('-')[0])
                e_y = int(end.split('-')[0])
                total_years += (e_y - s_y)
        except: pass
    
    # Extract top achievement
    achievements = []
    for emp in resume_json.get("employment_history", []):
        highlights = emp.get("highlights") or []
        if isinstance(highlights, list):
            achievements.extend([h for h in highlights if h])
    
    # Simple selection of top 3 achievements
    achievement_vals = achievements[:3]
    while len(achievement_vals) < 3:
        achievement_vals.append("N/A")
    
    # Project list
    projects = []
    for proj in resume_json.get("projects", []):
        p_name = proj.get("name")
        if p_name: projects.append(p_name)
    
    # Mocking/Extracting JD facts (In real app, this should come from a structured JD analysis)
    # For now, we'll use defaults or simple extraction
    facts = {
        "name": name,
        "email": person.get("email", "N/A"),
        "phone": person.get("phone", "N/A"),
        "location": person.get("location", "N/A"),
        "current_title": resume_json.get("employment_history", [{}])[0].get("title", "Professional"),
        "years_exp": f"{total_years}+",
        "achievement_1": achievement_vals[0],
        "achievement_2": achievement_vals[1],
        "achievement_3": achievement_vals[2],
        "project_list": ", ".join(projects) if projects else "N/A",
        "target_role": job_title or "Target Role",
        "target_company": company or "Target Company",
        "company_name": company or "Target Company", # Alias for user prompt
        "hiring_manager": "Hiring Team", # Default for now
        "target_location": "Remote / On-site", # Default
        "target_mission": "Building innovative solutions", # Default
        "target_keywords": "efficiency, scale, innovation", # Default
        "target_pain_points": "manual processes, scalability bottlenecks" # Default
    }
    
    return facts


def sanitize_job_title(title: str) -> str:
    """
    Remove annotations like (PROOF OF CONCEPT), (CONTRACT), (REMOTE) etc.
    and clean up whitespace.
    """
    if not title:
        return ""
    # Remove annotations in parentheses
    title = re.sub(r'\([^)]*\)', '', title)
    # Remove extra whitespace
    title = re.sub(r'\s+', ' ', title).strip()
    # Title Case for consistency if needed, though original casing is often preferred
    return title

def strip_prose_advice(text: str) -> str:
    """
    Remove common AI metacommentary and advice phrases.
    """
    if not text or not isinstance(text, str):
        return text or ""
        
    advice_patterns = [
        r"(?i)^Sure,?\s+here('s| is)\s+the\s+tailored\s+resume.*?\n",
        r"(?i)^Here\s+is\s+the\s+tailored\s+resume.*?\n",
        r"(?i)^I've\s+tailored\s+the\s+resume.*?\n",
        r"(?i)Hope\s+this\s+helps.*?\Z",
        r"(?i)Let\s+me\s+know\s+if\s+you\s+need\s+anything\s+else.*?\Z",
        r"(?i)PRD", # Strict removal of PRD text
    ]
    
    for pattern in advice_patterns:
        text = re.sub(pattern, "", text).strip()
    return text

async def generate_expert_tailored_content(resume_text: str, job_description: str, intensity_instruction: str = "") -> dict:
    """
    Expert Tailoring Pipeline (Phase 21):
    One-shot, high-fidelity generation of tailored resume data + cover letter.
    Using DeepSeek with detailed system instructions for maximum quality.
    """
    try:
        if not resume_text or len(str(resume_text)) < 50:
            logger.error(f"[Validation Guard] Resume content is missing or too short. Length: {len(str(resume_text)) if resume_text else 0}")
            return {"error": "Resume content is missing or too short. Could not load original resume data before tailoring."}
        
        logger.info(f"[Validation Guard] Sending Resume Text (length: {len(str(resume_text))}):\n{str(resume_text)[:100]}...")
        logger.info("Starting Expert AI Tailoring Pipeline...")
        
        # 1. Prepare User Prompt
        prompt = EXPERT_TAILORING_USER_PROMPT.format(
            raw_resume_text=resume_text,
            job_description=job_description
        )
        
        # 2. Call LLM with Expert System Prompt and specifically configured temperature
        response_raw = await unified_api_call(
            prompt, 
            max_tokens=8000, # Increased for detailed comprehensive output
            temperature=0.3, # Specific requested temperature for creative yet professional output
            system_prompt=EXPERT_SYSTEM_PROMPT,
            json_mode=True
        ) or ""
        
        # 3. Extract and parse JSON
        result = extract_json_from_response(response_raw)
        
        if not result or "tailored_resume" not in result:
            logger.error(f"Expert tailoring failed to return valid JSON. Raw result length: {len(response_raw)}")
            # Fallback attempt if JSON extraction failed but we have text
            if response_raw.strip().startswith('{'):
                 try:
                     result = json.loads(response_raw)
                 except: pass
            
            if not result or "tailored_resume" not in result:
                return {"error": "Expert tailoring failed to generate valid content"}
            
        # Recursive cleanup of any leaked advice or placeholders
        result["tailored_resume"] = recursive_cleanup(result["tailored_resume"])
        if "cover_letter" in result:
            result["cover_letter"] = strip_prose_advice(result["cover_letter"])
            
        logger.info("Expert AI Tailoring complete.")
        return result
        
    except Exception as e:
        logger.error(f"Expert tailoring error: {e}")
        return {"error": str(e)}


def strip_summary_garbage(text: str) -> str:
    """
    Remove known garbage fragments the AI copies from weak original summaries.
    Also strips generic filler openers that recruiters ignore.
    """
    if not text or not isinstance(text, str):
        return text or ""
        
    garbage_fragments = [
        r'\.\s*quality solutions\.?',
        r'\.\s*driven automation[,.]?',
        r',\s*quality solutions\.?',
        r'\.\s*governed,\s*documented data products\.?',
        r',\s*governed,\s*documented data products\.?',
        r'and drive business outcomes\.?',
        r'\s*I am excited (about the opportunity )?to apply my skills[^.]*\.',
        r'\s*particularly in a company like[^.]*\.',
        r'\s*that values innovation and excellence\.?',
        r'to the table[,.]?',
        r'Highly motivated\s+',
        r'looking to leverage my skills in a new domain[,.]?\s*',
        r'passionate about[^.]*\.',
        r'results-driven professional\s+',
        r'dynamic professional\s+',
        r'[\.\,]\s*oriented middleware[^.]*\.',     # "message-oriented middleware" fragment
        r'^[Oo]riented middleware[^.]*\.\s*',        # orphaned sentence start
        r'[\.\,]\s*functional teams to drive[^.]*\.',
        r'^[Ff]unctional teams to drive[^.]*\.\s*',
        r'\s*I am excited[^.]*\.',
        r'\s*particularly in a company like[^.]*\.',
        r'\s*that values innovation and excellence\.?',
        r'and drive business outcomes\.?',
        r'oriented middleware[,.]',
        r'\s*quality solutions that meet business outcomes\.?',
        r'\s*Excited to apply my skills[^.]*\.',
        r'\s*and contribute to the development of[^.]*\.',
        r"\s*Fanatics Betting\s*[&&]\s*Gaming[^.]*\.",
        r'\s*digital sports platform\.?',
        r'quality solutions[,.]?',
    ]
    for pattern in garbage_fragments:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # Clean up double spaces or double periods left behind
    text = re.sub(r'  +', ' ', text)
    text = re.sub(r'\.\s*\.', '.', text)
    return text.strip()





def format_resume_date(date_str: str) -> str:
    """Convert ISO date (2026-02) to Month Year (Feb 2026) for professional rendering."""
    if not date_str or str(date_str).lower() == "present":
        return "Present"
    import re
    # Handle ISO YYYY-MM
    m = re.match(r'^(\d{4})-(\d{2})$', str(date_str).strip())
    if m:
        year, month = m.groups()
        months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        try:
            idx = int(month) - 1
            if 0 <= idx < 12:
                return f"{months[idx]} {year}"
        except: pass
    # Handle YYYY-MM-DD
    m2 = re.match(r'^(\d{4})-(\d{2})-\d{2}$', str(date_str).strip())
    if m2:
        year, month = m2.groups()
        months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        try:
            idx = int(month) - 1
            if 0 <= idx < 12:
                return f"{months[idx]} {year}"
        except:
            pass
    return str(date_str)

def recursive_cleanup(data: Any) -> Any:
    """Recursively apply strip_prose_advice to all strings in a nested structure."""
    if isinstance(data, str):
        return strip_prose_advice(data)
    elif isinstance(data, list):
        return [recursive_cleanup(item) for item in data]
    elif isinstance(data, dict):
        return {k: recursive_cleanup(v) for k, v in data.items()}
    return data


def render_preview_text_from_json(data: Dict) -> str:
    """
    Converts the structured JSON from generate_optimized_resume_content 
    OR generate_expert_tailored_content into a plain text format that ResumePaper can parse.
    """
    if not data:
        return ""
        
    # Support for Phase 21 Expert Tailoring structure
    is_expert = isinstance(data, dict) and "tailored_resume" in data
    source = data.get("tailored_resume") if is_expert and isinstance(data, dict) else data
    
    if not isinstance(source, dict):
        # Fallback for when AI returns a string instead of JSON
        return str(source) if source else ""

    out = []
    
    # Header
    if is_expert:
        contact = source.get("contact", {}) if isinstance(source.get("contact"), dict) else {}
        fullName = str(source.get('name', 'Your Name')).strip()
        fullName = re.sub(r'(?i)^undefined\s*', '', fullName).strip()
        
        target_role = source.get("title") or ""
        if str(target_role).lower() in ["undefined", "none", "null"]:
            target_role = ""
            
        target_role = sanitize_job_title(target_role)
        
        # Final safety cleanup for fullName - aggressive to catch "UNDEFINEDSAIRAM"
        fullName = re.sub(r'(?i)undefined|none|null', '', fullName).strip()
        
        # Remove target_role from fullName if it's naming twice
        if target_role and fullName:
            # Strip target_role aggressively using partial matches
            # Split by - or | to get the base role
            role_base = re.split(r'[-|]', str(target_role))[0].strip()
            if len(role_base) > 3:
                pattern = re.escape(role_base)
                fullName = re.sub(f'(?i){pattern}', '', fullName).strip()
            
            # Also try the full target_role just in case
            full_pattern = re.escape(str(target_role))
            fullName = re.sub(f'(?i){full_pattern}', '', fullName).strip()
            
            # Final cleanup of any trailing separators usually accidentally left by AI
            fullName = re.sub(r'^[-\|\s]+|[-\|\s]+$', '', fullName).strip()
        
        if not fullName: fullName = "Your Name"
        
        out.append(f"NAME\n{fullName.upper()}")
        if target_role:
            out.append(target_role.upper())
    else:
        contact = source.get("contactInfo", {}) if isinstance(source.get("contactInfo"), dict) else {}
        fullName = str(contact.get('name', 'Your Name')).strip()
        fullName = re.sub(r'(?i)^undefined\s*', '', fullName).strip()
        out.append(f"NAME\n{fullName.upper()}")
    
    contacts = []
    if isinstance(contact, dict):
        contacts = [
            contact.get("email"),
            contact.get("phone"),
            contact.get("location"),
            contact.get("linkedin") if not is_expert else None, 
            contact.get("website") if not is_expert else None
        ]
    contact_line = " | ".join([c for c in contacts if c])
    out.append(f"CONTACT\n{contact_line}")
    
    # Summary
    summary = source.get('summary', '')
    if isinstance(summary, list):
        summary = " ".join(summary)
    summary = strip_prose_advice(summary)
    summary = strip_summary_garbage(summary)
    out.append(f"SUMMARY\n{summary}")
    
    # Skills
    skills = source.get("skills", [])
    if isinstance(skills, dict):
        skills_lines = []
        for cat, list_s in skills.items():
            skills_lines.append(f"• {cat}: {', '.join(list_s)}")
        out.append(f"SKILLS\n" + "\n".join(skills_lines))
    else:
        out.append(f"SKILLS\n" + "\n".join([f"• {s}" for s in skills]))
    
    # Experience
    out.append("EXPERIENCE")
    exp_entries = []
    for job in source.get("experience", []):
        company = job.get('company', '')
        title = job.get('title', '')
        location = job.get('location', '')
        dates = job.get('dates') or f"{job.get('start', '')} - {job.get('end', '')}".strip(" -")
        
        job_header = f"{company} — {title} | {location}\n{dates}"
        entry_text = f"{job_header}\n"
        
        if job.get("project_summary"):
            entry_text += f"\nProject Summary: {job['project_summary']}\n"
            entry_text += "\nRoles and Responsibilities:\n"
            
        bullets = "\n".join([f"- {b}" for b in job.get("bullets", [])])
        entry_text += bullets
        
        if job.get("environment"):
            entry_text += f"\n\nEnvironment: {job['environment']}"
            
        exp_entries.append(entry_text)
    out.append("\n".join(exp_entries))
            
    # Projects
    if source.get("projects"):
        out.append("PROJECTS")
        proj_entries = []
        for proj in source.get("projects", []):
            name = proj.get('name', '')
            subtitle = proj.get('subtitle', '') or proj.get('tech', '')
            proj_header = f"{name} — {subtitle}"
            bullets = "\n".join([f"- {b}" for b in proj.get("bullets", [])])
            proj_entries.append(f"{proj_header}\n{bullets}")
        out.append("\n".join(proj_entries))
                
    # Certifications
    if source.get("certifications"):
        out.append("CERTIFICATIONS")
        cert_lines = []
        for cert in source.get("certifications", []):
            if isinstance(cert, str):
                cert_lines.append(f"• {cert}")
            elif isinstance(cert, dict):
                cert_lines.append(f"• {cert.get('name', '')}")
        out.append("\n".join(cert_lines))
                
    # Education
    if source.get("education"):
        out.append("EDUCATION")
        edu_lines = []
        for edu in source.get("education", []):
            edu_line = f"{edu.get('degree', '')}, {edu.get('school', '') or edu.get('university', '')} | {edu.get('date', '') or edu.get('year', '')}"
            edu_lines.append(edu_line)
        out.append("\n".join(edu_lines))
            
    return "\n\n".join(out)


def _render_ats_with_dynamic_skills(r: 'ResumeDataSchema', skill_labels: List[str]) -> str:
    """
    Renders the ATS resume using the AI-returned skill category labels
    instead of the hardcoded CoreSkills field names.
    """
    import re
    header_parts = [
        r.header.city_state,
        r.header.phone,
        r.header.email,
        r.header.linkedin,
        r.header.portfolio
    ]
    header_line = " | ".join([p for p in header_parts if p and str(p).strip() and str(p).lower() != "undefined"])
    header_line = re.sub(r'(?i)undefined', '', header_line)

    out = []
    # Bug 1 Fix: Preventive name/role building
    name = (r.header.full_name or "").strip()
    target_role = (r.header.target_role or "").strip()
    
    # Clean undefined leaks
    name = re.sub(r'(?i)undefined|none|null', '', name).strip()
    target_role = re.sub(r'(?i)undefined|none|null', '', target_role).strip()
    
    target_role = sanitize_job_title(target_role)
    
    if not name:
        name = "Your Name"
        
    out.append(name.upper())
    if target_role:
        out.append(target_role.upper())
    out.append(header_line)
    out.append("")

    out.append("PROFESSIONAL SUMMARY")
    summary_text = r.positioning_statement
    summary_text = summary_text.replace('\\n', '\n').replace('\\\\n', '\n')
    out.append(summary_text)
    out.append("")

    out.append("SKILLS")
    # Use the AI-returned category labels (matched by index to CoreSkills fields)
    skill_fields = [
        r.core_skills.languages,
        r.core_skills.data_etl,
        r.core_skills.cloud,
        r.core_skills.databases,
        r.core_skills.devops_tools,
        r.core_skills.other,
    ]
    for i, label in enumerate(skill_labels):
        if i < len(skill_fields) and skill_fields[i]:
            out.append(f"{label}: {', '.join(skill_fields[i])}")
    out.append("")

    if r.experience:
        out.append("EXPERIENCE")
        for role in r.experience:
            company = (role.company or "").strip()
            title = (role.job_title or "").strip()
            location = (role.city_state_or_remote or "").strip()
            
            # Professional Header: COMPANY | TITLE | DATE
            start_f = format_resume_date(role.start)
            end_f = format_resume_date(role.end)
            dates_str = f"{start_f} – {end_f}"
            
            header_parts = [company, title, dates_str]
            out.append(" | ".join([p for p in header_parts if p]))
            if location:
                out.append(location)
                
            for b in role.bullets:
                out.append(f"• {cleanup_bullet(b)}")
            out.append("")

    if r.projects:
        out.append("PROJECTS")
        for p in r.projects:
            name = (p.name or "").strip()
            tech = ", ".join(p.tech_stack) if p.tech_stack else ""
            link = f" ({p.link})" if p.link else ""
            if name and tech:
                out.append(f"{name} — {tech}{link}")
            elif name:
                out.append(f"{name}{link}")
            elif tech:
                out.append(tech)
            for b in p.bullets:
                out.append(f"• {cleanup_bullet(b)}")
    if hasattr(r, 'certifications') and r.certifications:
        out.append("CERTIFICATIONS")
        for c in r.certifications:
            out.append(f"• {c}")
        out.append("")

    if r.education:
        out.append("EDUCATION")
        for e in r.education:
            # Bug 4 Fix: Degree in Major | University | Year
            degree_part = f"{e.degree} in {e.major}" if e.degree and e.major else (e.degree or e.major)
            if degree_part:
                out.append(degree_part)
            
            uni_year = " | ".join(filter(None, [e.university, e.year]))
            if uni_year:
                out.append(uni_year)
            out.append("")

    return "\n".join(out).strip()


def render_ats_resume_from_json(r: ResumeDataSchema) -> str:
    """
    Deterministic rendering of the ATS template (Step 4).
    Guarantees same headings, order, and style.
    """
    header_parts = [
        r.header.city_state,
        r.header.phone,
        r.header.email,
        r.header.linkedin,
        r.header.portfolio
    ]
    header_line = " | ".join([p for p in header_parts if p and str(p).strip() and str(p).lower() != "undefined"])
    import re
    header_line = re.sub(r'(?i)undefined', '', header_line)

    out = []
    # Bug 1 Fix: Preventive name/role building
    name = (r.header.full_name or "").strip()
    target_role = (r.target_title or "").strip()
    
    # Clean undefined leaks
    name = re.sub(r'(?i)undefined|none|null', '', name).strip()
    target_role = re.sub(r'(?i)undefined|none|null', '', target_role).strip()
    
    if not name:
        name = "Your Name"

    header_block = f"{name.upper()}\n{target_role.upper()}" if target_role else name.upper()
    out.append(header_block)
    out.append(header_line)
    out.append("")

    out.append("PROFESSIONAL SUMMARY")
    summary_text = r.positioning_statement
    summary_text = summary_text.replace('\\n', '\n').replace('\\\\n', '\n')
    out.append(summary_text)
    out.append("")

    out.append("SKILLS")
    if r.core_skills:
        skill_categories = [
            ("Languages", r.core_skills.languages),
            ("Data/ETL", r.core_skills.data_etl),
            ("Cloud", r.core_skills.cloud),
            ("Databases", r.core_skills.databases),
            ("DevOps/Tools", r.core_skills.devops_tools),
            ("Other", r.core_skills.other),
        ]
        count = 0
        for cat_name, cat_list in skill_categories:
            if cat_list:
                out.append(f"{cat_name}: {', '.join(cat_list)}")
                count += 1
    out.append("")

    if r.experience:
        out.append("EXPERIENCE")
        for role in r.experience:
            company = (role.company or "").strip()
            title = (role.job_title or "").strip()
            location = (role.city_state_or_remote or "").strip()
            
            # Header: COMPANY | TITLE | DATE
            start_f = format_resume_date(role.start)
            end_f = format_resume_date(role.end)
            dates_str = f"{start_f} – {end_f}"
            
            header_parts = [company, title, dates_str]
            out.append(" | ".join([p for p in header_parts if p]))
            if location:
                out.append(location)
                
            for b in role.bullets:
                out.append(f"• {cleanup_bullet(b)}")
            out.append("")

    if r.projects:
        out.append("PROJECTS")
        for p in r.projects:
            name = (p.name or "").strip()
            tech = ", ".join(p.tech_stack) if p.tech_stack else ""
            link = f" ({p.link})" if p.link else ""
            if name and tech:
                out.append(f"{name} — {tech}{link}")
            elif name:
                out.append(f"{name}{link}")
            elif tech:
                out.append(tech)
            for b in p.bullets:
                out.append(f"• {cleanup_bullet(b)}")
            out.append("")

    if r.certifications:
        out.append("CERTIFICATIONS")
        for c in r.certifications:
            out.append(f"• {c}")
        out.append("")

    if r.education:
        out.append("EDUCATION")
        for e in r.education:
            degree_part = " ".join(filter(None, [e.degree, e.major]))
            if degree_part:
                out.append(degree_part)
            
            uni_year = " | ".join(filter(None, [e.university, e.year]))
            if uni_year:
                out.append(uni_year)
            out.append("")

    return "\n".join(out).strip()


async def generate_simple_tailored_resume(resume_text: str, job_description: str, job_title: str, company: str, missing_skills_list: list = None) -> str:
    """Modular 3-Stage Resume Transformation: Analysis -> Granular Rewriting -> Assembly"""
    try:
        logger.info(f"Starting 3-stage tailoring pipeline for {company}")
        missing_skills = json.dumps(missing_skills_list or [])
        
        # 0. Extract Structured Data
        resume_json = await extract_resume_data(resume_text)
        if not resume_json or "error" in resume_json:
            logger.warning("Failed to extract structured resume data. Falling back to single-shot.")
            return await _generate_single_shot_fallback(resume_text, job_description, job_title, company)

        # STAGE 1: ANALYSIS
        analysis_prompt = f"""
You are a resume analyst. Output ONLY valid JSON. No explanation, no markdown, no backticks.

Given the resume and job description, return this exact JSON structure:

{{
  "jd_title": "exact job title from JD",
  "jd_company": "company name",
  "jd_domain": "one sentence describing company mission",
  "jd_seniority": "Mid / Senior / etc",
  "jd_required_skills": ["skill1", "skill2", "skill3", "skill4", "skill5", "skill6", "skill7", "skill8"],
  "candidate_has": ["skills from JD that exist in candidate resume"],
  "candidate_lacks": ["skills from JD absent from candidate resume"],
  "missing_skills_routing": {{
    "SkillName": "skills_section_only"
  }},
  "title_reframes": {{
    "Employer Name 1": "reframed title toward JD",
    "Employer Name 2": "reframed title toward JD"
  }},
  "summary_keywords": ["5 JD keywords candidate actually has to use in summary"],
  "top_2_projects": ["most JD-relevant project name 1", "most JD-relevant project name 2"]
}}

RESUME:
{json.dumps(resume_json, indent=2)}

JOB DESCRIPTION:
{job_description}

MISSING SKILLS SELECTED BY USER:
{missing_skills}
"""
        analysis_resp = await unified_api_call(analysis_prompt, max_tokens=1000, model="llama-3.1-8b-instant", json_mode=True)
        analysis_data = json.loads(clean_json_response(analysis_resp))
        logger.info(f"Stage 1 Analysis complete for {company}")

        # STAGE 2: ROLE REWRITING (Granular)
        rewritten_roles = []
        for role in resume_json.get("employment_history", []):
            role_company = role.get("company", "Unknown")
            original_title = role.get("title", "Unknown")
            # Handle plural title_reframes
            new_title = analysis_data.get("title_reframes", {}).get(role_company, original_title)
            
            rewrite_prompt = f"""
You are a professional resume writer. Rewrite ONE job role's bullets for a resume.
Respond with ONLY a JSON object. No prose.

CONTEXT:
- Company: {role_company}
- Original title: {original_title}  
- Reframed title: {new_title}
- Original bullets: {json.dumps(role.get('highlights', role.get('description', [])), indent=2)}
- Skills candidate ACTUALLY HAS: {json.dumps(analysis_data.get('candidate_has', []), indent=2)}

OUTPUT SCHEMA:
{{ "bullets": ["bullet 1", "bullet 2", "bullet 3", "bullet 4", "bullet 5"] }}
"""
            rewrite_resp_raw = await unified_api_call(rewrite_prompt, max_tokens=12000, model="llama-3.1-8b-instant", json_mode=True) or ""
            rewrite_data = extract_json_from_response(rewrite_resp_raw)
            bullets = rewrite_data.get("bullets", [])
            bullets_text = "\n".join([f"• {b}" for b in bullets if b])
            
            rewritten_roles.append(f"{role_company} | {role.get('startDate')} - {role.get('endDate')}\n{new_title} | {role.get('location')}\n{bullets_text.strip()}")

        logger.info(f"Stage 2 Rewriting complete for {len(rewritten_roles)} roles")

        assembly_prompt = f"""
You are a professional resume writer. Respond ONLY with a JSON object. No prose.

{json.dumps({
    "analysis": analysis_data,
    "rewritten_roles": rewritten_roles,
    "missing_skills": missing_skills,
    "original_resume": resume_json
})}

---
WRITE EACH SECTION IN THE 'resume_text' FIELD AS FOLLOWS:

1. PROFESSIONAL SUMMARY: 3-4 sentences of natural flowing prose. Start with "[X]+ years of experience".
2. SKILLS: 4 labeled categories.
3. EXPERIENCE: Use the REWRITTEN ROLES provided in the input object.
4. PROJECTS: Write all projects with tech stack and 3 bullets each.
5. CERTIFICATIONS & EDUCATION: Keep from original.

OUTPUT SCHEMA:
{{ "resume_text": "FULL TAILORED RESUME TEXT HERE" }}
"""

        final_resp = await unified_api_call(
            assembly_prompt, max_tokens=4000, 
            model="llama-3.3-70b-versatile", json_mode=True
        )
        final_data = extract_json_from_response(final_resp)
        final_assembly = final_data.get("resume_text", "")
        
        # STAGE 4: Remove entirely — it causes advice-text output for mismatched roles
        # The assembly output is already clean enough
        final_resume = final_assembly

        # Final Cleanup
        final_resume = re.sub(r'\n{3,}', '\n\n', final_resume.strip())
        return final_resume

    except Exception as e:
        logger.error(f"Modular tailoring failed: {e}")
        return await _generate_single_shot_fallback(resume_text, job_description, job_title, company)

async def _generate_single_shot_fallback(resume_text: str, job_description: str, job_title: str, company: str) -> str:
    """Single-shot fallback logic if modular pipeline fails"""
    prompt = f"""CRITICAL: You are an automated resume tailoring engine. 
Output ONLY a JSON object. No prose. No markdown.

Tailor this resume for {job_title} at {company}:

RESUME: {resume_text}
JD: {job_description}

---
WRITE THE TAILORED RESUME IN THE 'resume_text' FIELD AS FOLLOWS:
1. SUMMARY: 3-4 sentences targeting the JD.
2. SKILLS: Optimized for ATS matching.
3. EXPERIENCE: Quantify impact and use JD keywords.
4. PROJECTS: Include all projects with tech stack.

OUTPUT SCHEMA:
{{ "resume_text": "FULL TAILORED RESUME TEXT HERE" }}
"""
    try:
        resp = await unified_api_call(prompt, max_tokens=3000, json_mode=True)
        data = extract_json_from_response(resp)
        return data.get("resume_text") or resume_text
    except:
        return resume_text


async def generate_optimized_resume_content(resume_text: str, job_description: str, analysis: Dict, target_score: int = 85) -> Optional[Dict]:
    """Generate optimized resume content using AI - achieves target ATS score while preserving authenticity"""
    
    missing_skills = []
    if analysis.get("hardSkills", {}).get("missing"):
        missing_skills.extend([s.get("skill", s) if isinstance(s, dict) else s for s in analysis["hardSkills"]["missing"]])
    if analysis.get("softSkills", {}).get("missing"):
        missing_skills.extend([s.get("skill", s) if isinstance(s, dict) else s for s in analysis["softSkills"]["missing"]])
    
    keywords = analysis.get("keywordsToAdd", [])
    matched_skills = []
    if analysis.get("hardSkills", {}).get("matched"):
        matched_skills.extend([s.get("skill", s) if isinstance(s, dict) else s for s in analysis["hardSkills"]["matched"]])
    
    prompt = f"""
You are an Elite Resume Architect and ATS Optimization Expert. Your goal is to rewrite/enhance the resume to achieve a {target_score}% ATS match score.

=== ABSOLUTE RULES - VIOLATION = FAILURE ===
1. EVERY job position in the original MUST appear in output (count them!)
2. EVERY bullet point MUST be preserved - you can ONLY add more or enhance existing ones, never remove
3. EVERY project MUST be included with ALL its details
4. Output MUST be professional, high-signal, and achieve {target_score}% ATS compatibility
5. Copy the EXACT job titles, company names, dates - do NOT paraphrase
6. Keep the original professional summary structure but enhance it with keywords

=== YOUR TASK ===
1. Optimize every bullet point for Impact and Results.
2. Use powerful action verbs (e.g., "orchestrated", "pioneered").
3. Add specific metrics/numbers where logical.
4. DO NOT append skill suffixes to bullets. Skills should be grouped in the Skills section ONLY.
5. Ensure content is industry-standard for: {job_description[:500]}

=== ORIGINAL RESUME (COPY ALL CONTENT - DO NOT SKIP ANY SECTION) ===
{resume_text}

=== TARGET JOB DESCRIPTION ===
{job_description}

=== OUTPUT FORMAT (JSON) ===
Return ONLY valid JSON, no markdown:

{{
    "contactInfo": {{
        "name": "COPY EXACT NAME",
        "email": "COPY EXACT EMAIL",
        "phone": "COPY EXACT PHONE",
        "location": "COPY EXACT LOCATION",
        "linkedin": "COPY IF EXISTS",
        "website": "COPY IF EXISTS"
    }},
    "summary": "ENHANCE the original summary to be 3-5 lines dense with high-impact keywords and your core value prop. Target {target_score}% alignment.",
    "experience": [
        {{
            "title": "COPY EXACT TITLE",
            "company": "COPY EXACT COMPANY NAME",
            "location": "COPY EXACT LOCATION",
            "dates": "COPY EXACT DATES",
            "bullets": [
                "ENHANCED version of original bullet 1 - integrate keywords and metrics",
                "ENHANCED version of original bullet 2",
                "ENHANCED version of original bullet 3",
                "... INCLUDE AND ENHANCE EVERY ORIGINAL BULLET ...",
                "ADD 1-2 extra high-impact bullets focusing on market-standard skills missed in original"
            ]
        }}
    ],
    "projects": [
        {{
            "name": "COPY EXACT PROJECT NAME",
            "subtitle": "Brief description or tech stack - enhance with keywords",
            "bullets": [
                "COPY and enhance all project details with impact metrics"
            ]
        }}
    ],
    "education": [
        {{
            "degree": "COPY EXACT DEGREE",
            "school": "COPY EXACT SCHOOL NAME",
            "location": "COPY IF EXISTS",
            "date": "COPY DATE"
        }}
    ],
    "skills": ["COPY ALL ORIGINAL SKILLS", "ADD market-standard technical and soft skills to reach {target_score}% match"]
}}

=== VERIFICATION CHECKLIST ===
- All positions included
- All bullets enhanced and preserved
- {target_score}% ATS optimization target met
- Skills are categorized and routed to the SKILLS section only
- Metrics added to at least 50% of bullets

GENERATE THE {target_score}% OPTIMIZED RESUME JSON NOW:
"""

    try:
        # Use unified call with JSON Mode enabled
        response = await unified_api_call(prompt, max_tokens=8000, model="llama-3.1-8b-instant", json_mode=True)
        if not response:
            return None
        
        # Robust cleaning using shared utility
        cleaned_json = clean_json_response(response)
        
        return json.loads(cleaned_json)
    except Exception as e:
        logger.error(f"Failed to generate resume content: {e}")
        return None


async def generate_cover_letter_content(resume_text: str, job_description: str, job_title: str, company: str) -> Optional[str]:
    """Generate a cover letter using AI"""
    
    prompt = f"""
Write a professional, compelling cover letter for this job application.

APPLICANT'S RESUME:
{resume_text[:2500]}

JOB TITLE: {job_title}
COMPANY: {company}

JOB DESCRIPTION:
{job_description[:2000]}

Write a 3-4 paragraph cover letter that:
1. Opens with enthusiasm for the specific role and company
2. Highlights 2-3 most relevant experiences/achievements from the resume
3. Shows knowledge of the company/industry
4. Closes with a strong call to action

Use a professional but personable tone. Do NOT use brackets or placeholders.
Return ONLY the cover letter text. Do NOT include any introductory or concluding advice, guides, or "Here is your cover letter". 
Start with the date or "Dear [Hiring Manager]," immediately.
"""

    try:
        response = await unified_api_call(prompt)
        return response
    except Exception as e:
        logger.error(f"Failed to generate cover letter: {e}")
        return None


async def extract_compliance_facts(resume_text: str) -> Optional[Dict]:
    """Stage 1: Extract verbatim facts from resume into a strict JSON schema"""
    # Truncate resume text only if massive
    truncated_resume = resume_text[:15000]
    
    prompt = f"""
SYSTEM:
You are a precision fact extractor. Output JSON ONLY. No prose.

USER:
Extract EVERY detail verbatim from the candidate resume below. 

ABSOLUTE RULES:
1. DO NOT SKIP ANY JOBS. 
2. DO NOT SKIP ANY BULLET POINTS.
3. DO NOT SUMMARIZE. Copy text exactly as it appears.
4. Output ONLY valid JSON.

[CANDIDATE_BASE_RESUME]
<<<
{truncated_resume}
>>>

Return JSON with this exact schema:
{{
  "name": "",
  "location": "",
  "email": "",
  "phone": "",
  "links": {{"linkedin":"", "github":"", "portfolio":""}},
  "summary_original": "",
  "employers": [{{"company":"","title":"","location":"","start":"","end":"","bullets":[] }}],
  "projects": [{{"name":"","bullets":[],"tools":[],"metrics":[] }}],
  "education": [{{"degree":"","major":"","university":"","year":""}}],
  "certifications": [],
  "skills": {{"technical":[], "soft":[]}},
  "metrics_explicit": []
}}

Respond with a valid JSON object only. Begin with {{ now:
"""
    try:
        # Use 70B for extraction to ensure NO content loss (Stage 1 High-Capacity Mode)
        response_text = await unified_api_call(prompt, max_tokens=5000, model="llama-3.3-70b-versatile", json_mode=True)
        if not response_text:
            return {}
        
        json_text = clean_json_response(response_text)
        result = json.loads(json_text)
        
        # DEBUG — remove after confirming output
        logger.info(f"Stage 1 facts_json employers: {[e.get('company','') + ' | ' + e.get('title','') for e in result.get('employers', [])]}")
        logger.info(f"Stage 1 facts_json projects: {[p.get('name','') for p in result.get('projects', [])]}")
        logger.info(f"Stage 1 facts_json education: {result.get('education', [])}")
        logger.info(f"Stage 1 facts_json certifications: {result.get('certifications', result.get('skills', {}).get('certifications', []))}")
        
        return result
    except Exception as e:
        logger.error(f"Fact extraction failed: {e}")
        return {}

async def generate_expert_documents(
    resume_text: str, 
    job_description: str, 
    user_info: Optional[Dict] = None, 
    selected_sections: Optional[List[str]] = None,
    selected_keywords: Optional[List[str]] = None,
    job_title: str = "",
    company: str = "",
    intensity: str = "default",  # low, default, high
    length_target: str = "standard", # standard, compact
    force_metrics: bool = False
) -> Optional[Dict[str, Any]]:
    """Generate ATS Resume and Detailed CV using the Expert AI Ninja Engine (Phase 21)"""
    
    try:
        if not resume_text or len(str(resume_text)) < 50:
            logger.error(f"[Validation Guard] Resume content is missing or too short. Length: {len(str(resume_text)) if resume_text else 0}")
            return {
                "alignment_highlights": [],
                "ats_resume": "",
                "detailed_cv": "",
                "cover_letter": "",
                "resume_json": {},
                "error": "Resume content is missing or too short. Could not load original resume data before tailoring."
            }
            
        # Intensity mapping
        if intensity == "aggressive":
            intensity_instruction = "RADICAL TAILORING: Rewrite bullets to use as many JD keywords as possible while maintaining factual truth. Prioritize impact and matching."
        elif intensity == "minimal":
            intensity_instruction = "MINIMAL TAILORING: Keep original bullets mostly intact, only adding 1-2 keywords where they naturally fit."
        else:
            intensity_instruction = "BALANCED TAILORING: Professional reframing that aligns candidate strengths with JD requirements."

        # 1. Extract structured facts for documents
        resume_json = await extract_resume_data(resume_text)
        
        # 2. Generate Expert Tailored content (Resume)
        expert_data = await generate_expert_tailored_content(
            resume_text, job_description, intensity_instruction=intensity_instruction
        )
        
        # 3. Generate Cover Letters and Cold Emails using the new prompt
        cl_ce_facts = prepare_cover_letter_cold_email_facts(resume_json, job_description, job_title, company)
        cl_ce_prompt = COVER_LETTER_AND_COLD_EMAIL_PROMPT.format(**cl_ce_facts)
        
        cl_ce_resp = await unified_api_call(
            cl_ce_prompt,
            max_tokens=4000,
            model="llama-3.3-70b-versatile",
            json_mode=True
        )
        cl_ce_data = extract_json_from_response(cl_ce_resp)
        if not isinstance(cl_ce_data, dict): cl_ce_data = {}
        
        if not expert_data or not isinstance(expert_data, dict) or "error" in expert_data:
            err_msg = expert_data.get("error") if isinstance(expert_data, dict) else "Unknown error"
            logger.error(f"Expert generation failed: {err_msg}")
            # Fallback to simple tailoring if expert fails
            simple_resume = await generate_simple_tailored_resume(resume_text, job_description, job_title, company)
            cover_letter = await generate_cover_letter_content(resume_text, job_description, job_title, company)
            return {
                "alignment_highlights": [],
                "ats_resume": simple_resume,
                "detailed_cv": simple_resume,
                "cover_letter": cover_letter,
                "resume_json": {},
                "changes": [],
                "skills_added": [],
                "skills_skipped": []
            }

        # Format successfully
        tailored_res = expert_data.get("tailored_resume", {})
        rendered_text = render_preview_text_from_json(expert_data)
        
        jd_analysis = expert_data.get("step2_jd_analysis", {})
        if not isinstance(jd_analysis, dict): jd_analysis = {}

        return {
            "alignment_highlights": jd_analysis.get("must_have", []),
            "ats_resume": rendered_text,
            "detailed_cv": rendered_text,
            "cover_letter": cl_ce_data.get("cover_letter_A", expert_data.get("cover_letter", "")),
            "cover_letter_A": cl_ce_data.get("cover_letter_A", ""),
            "cover_letter_B": cl_ce_data.get("cover_letter_B", ""),
            "cold_email_A": cl_ce_data.get("cold_email_A", ""),
            "cold_email_B": cl_ce_data.get("cold_email_B", ""),
            "resume_json": tailored_res,
            "changes": [],
            "skills_added": [], # Timeline law is internal to prompt now
            "skills_skipped": []
        }
    except Exception as e:
        logger.error(f"Failed to generate expert documents: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        
        # Comprehensive fallback
        try:
            logger.info(f"Expert generation failed. Falling back to simple tailoring for {job_title} at {company}...")
            simple_text = await generate_simple_tailored_resume(
                resume_text, job_description, job_title or "Position", company or "Company",
                missing_skills_list=selected_keywords
            )
            cl_text = await generate_cover_letter_content(resume_text, job_description, job_title or "Position", company or "Company")
            simple_text = strip_appended_skill_suffixes(simple_text)
            simple_text = strip_summary_prefix(simple_text)
            simple_text = strip_summary_garbage(simple_text)
            return {
                "alignment_highlights": ["- Tailored analysis complete (fallback mode)"],
                "ats_resume": simple_text, 
                "detailed_cv": simple_text,
                "cover_letter": cl_text, 
                "cover_letter_A": cl_text,
                "cover_letter_B": "",
                "cold_email_A": "",
                "cold_email_B": "",
                "resume_json": {},
                "changes": [], 
                "skills_added": [], 
                "skills_skipped": []
            }
        except Exception as fallback_err:
            logger.error(f"Critical Tailoring Failure (Primary AND Fallback failed): {fallback_err}")
            return None

    return None

def create_resume_docx(resume_data: Dict, font_family: str = "Times New Roman", font_size: int = 11) -> io.BytesIO:
    """Create a comprehensive Word document from resume data"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_family
    font.size = Pt(font_size)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Contact Info / Header
    contact = resume_data.get("contactInfo", {})
    name = contact.get("name", "Your Name")
    
    # Name - Large and Bold
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(name.upper()) # Force upper as requested
    name_run.bold = True
    name_run.font.size = Pt(14) # Pt(14) corresponds to user's "size=28" in some contexts, but library uses Pt
    name_run.font.name = "Arial"
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    
    # Target Role - Smaller, Centered, Not Bold
    target_role = sanitize_job_title(contact.get("target_role") or contact.get("title") or "")
    if target_role:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(target_role.upper())
        title_run.bold = False
        title_run.font.size = Pt(11)
        title_run.font.name = "Arial"
        title_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(12)
    
    # Contact details - First line (email, phone, location)
    contact_line1 = []
    if contact.get("email"):
        contact_line1.append(contact["email"])
    if contact.get("phone"):
        contact_line1.append(contact["phone"])
    if contact.get("location"):
        contact_line1.append(contact["location"])
    
    if contact_line1:
        contact_para1 = doc.add_paragraph(" | ".join(contact_line1))
        contact_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para1.paragraph_format.space_after = Pt(0)
    
    # Contact details - Second line (linkedin, website)
    contact_line2 = []
    if contact.get("linkedin"):
        contact_line2.append(contact["linkedin"])
    if contact.get("website"):
        contact_line2.append(contact["website"])
    
    if contact_line2:
        contact_para2 = doc.add_paragraph(" | ".join(contact_line2))
        contact_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para2.paragraph_format.space_after = Pt(2)
    
    # Professional Summary
    if resume_data.get("summary"):
        heading = doc.add_heading("PROFESSIONAL SUMMARY", level=1)
        heading.runs[0].font.size = Pt(12)
        summary_para = doc.add_paragraph(resume_data["summary"])
        summary_para.paragraph_format.space_after = Pt(2)
    
    # Professional Experience
    if resume_data.get("experience"):
        heading = doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for i, exp in enumerate(resume_data["experience"]):
            # Title and Company on same line
            exp_header = doc.add_paragraph()
            title_run = exp_header.add_run(f"{exp.get('title', 'Title')}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            company_text = f" | {exp.get('company', 'Company')}"
            if exp.get("location"):
                company_text += f", {exp['location']}"
            exp_header.add_run(company_text)
            exp_header.paragraph_format.space_after = Pt(0)
            
            # Project Summary
            if exp.get("project_summary"):
                proj_sum_para = doc.add_paragraph()
                proj_sum_run = proj_sum_para.add_run(f"Project Summary: {exp['project_summary']}")
                proj_sum_para.paragraph_format.space_after = Pt(2)

            # Roles and Responsibilities Heading
            rar_para = doc.add_paragraph()
            rar_run = rar_para.add_run("Roles and Responsibilities:")
            rar_run.bold = True
            rar_para.paragraph_format.space_after = Pt(0)

            # Bullets - all of them
            bullets = exp.get("bullets", [])
            for bullet in bullets:
                if bullet and bullet.strip():
                    # Check if bullet contains bold markers
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    
                    # Simple bold logic for markdown-like **bold** in bullets
                    text = bullet.strip()
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = bullet_para.add_run(part[2:-2])
                            run.bold = True
                        else:
                            bullet_para.add_run(part)
                    
                    bullet_para.paragraph_format.space_after = Pt(0)
            
            # Environment
            if exp.get("environment"):
                env_para = doc.add_paragraph()
                env_label = env_para.add_run("Environment: ")
                env_label.bold = True
                env_para.add_run(exp["environment"])
                env_para.paragraph_format.space_before = Pt(2)
                env_para.paragraph_format.space_after = Pt(6)
            
            # Add space between experiences
            if i < len(resume_data["experience"]) - 1:
                pass # Removed extra spacer between roles
    
    # Projects Section
    if resume_data.get("projects"):
        heading = doc.add_heading("PROJECTS", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for project in resume_data["projects"]:
            proj_header = doc.add_paragraph()
            name_run = proj_header.add_run(f"{project.get('name', 'Project')}")
            name_run.bold = True
            
            # Handle subtitle or technologies
            subtitle = project.get("subtitle") or project.get("technologies")
            if subtitle:
                proj_header.add_run(f" — {subtitle}")
            proj_header.paragraph_format.space_after = Pt(0)
            
            if project.get("description"):
                desc_para = doc.add_paragraph(project["description"])
                desc_para.paragraph_format.space_after = Pt(2)
            
            # Handle bullets - could be list or single string
            bullets = project.get("bullets", [])
            if isinstance(bullets, str):
                bullets = [bullets]
            for bullet in bullets:
                if bullet and bullet.strip():
                    # Check if it's an "Impact:" section
                    if bullet.lower().startswith("impact:"):
                        impact_para = doc.add_paragraph()
                        impact_label = impact_para.add_run("Impact: ")
                        impact_label.bold = True
                        impact_para.add_run(bullet[7:].strip())
                    else:
                        doc.add_paragraph(bullet, style='List Bullet')
    
    # Education
    if resume_data.get("education"):
        heading = doc.add_heading("EDUCATION", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for edu in resume_data["education"]:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(f"{edu.get('degree', 'Degree')}")
            degree_run.bold = True
            edu_para.add_run(f" | {edu.get('school', 'School')}")
            if edu.get("location"):
                edu_para.add_run(f", {edu['location']}")
            edu_para.paragraph_format.space_after = Pt(0)
            
            # Date and GPA
            details_parts = []
            if edu.get("date"):
                details_parts.append(edu["date"])
            if edu.get("gpa"):
                details_parts.append(f"GPA: {edu['gpa']}")
            if details_parts:
                details_para = doc.add_paragraph(" | ".join(details_parts))
                details_para.runs[0].italic = True
                details_para.paragraph_format.space_after = Pt(2)
            
            if edu.get("details"):
                doc.add_paragraph(edu["details"])
    
    # Skills - Handle both list and dict formats
    skills_data = resume_data.get("skills", {})
    if skills_data:
        heading = doc.add_heading("SKILLS", level=1)
        heading.runs[0].font.size = Pt(12)
        
        if isinstance(skills_data, dict):
            # New format with categories
            if skills_data.get("technical"):
                tech_para = doc.add_paragraph()
                tech_label = tech_para.add_run("Technical: ")
                tech_label.bold = True
                tech_para.add_run(" • ".join(skills_data["technical"]))
            
            if skills_data.get("tools"):
                tools_para = doc.add_paragraph()
                tools_label = tools_para.add_run("Tools & Platforms: ")
                tools_label.bold = True
                tools_para.add_run(" • ".join(skills_data["tools"]))
            
            if skills_data.get("other"):
                other_para = doc.add_paragraph()
                other_label = other_para.add_run("Other: ")
                other_label.bold = True
                other_para.add_run(" • ".join(skills_data["other"]))
        else:
            # Old format - simple list
            skills_text = " • ".join(skills_data)
            doc.add_paragraph(skills_text)
    
    # Certifications
    if resume_data.get("certifications"):
        heading = doc.add_heading("CERTIFICATIONS", level=1)
        heading.runs[0].font.size = Pt(12)
        
        for cert in resume_data["certifications"]:
            cert_para = doc.add_paragraph()
            cert_name = cert_para.add_run(f"{cert.get('name', 'Certification')}")
            cert_name.bold = True
            if cert.get("issuer"):
                cert_para.add_run(f" - {cert['issuer']}")
            if cert.get("date"):
                cert_para.add_run(f" ({cert['date']})")
    
    # Additional Sections (Awards, Publications, Volunteer, etc.)
    if resume_data.get("additionalSections"):
        for section in resume_data["additionalSections"]:
            if section.get("title") and section.get("items"):
                heading = doc.add_heading(section["title"].upper(), level=1)
                heading.runs[0].font.size = Pt(12)
                for item in section["items"]:
                    if item and item.strip():
                        doc.add_paragraph(f"• {item}")
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def create_cover_letter_docx(cover_letter_text: str, job_title: str, company: str, font_family: str = "Times New Roman") -> io.BytesIO:
    """Create a Word document for the cover letter"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_family
    font.size = Pt(11)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Date
    from datetime import datetime
    date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Spacing
    
    # Greeting
    doc.add_paragraph(f"Re: Application for {job_title} at {company}")
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph()
    
    # Body - split into paragraphs
    # Process text
    text = (cover_letter_text or "").replace('\\n', '\n').replace('\\\\n', '\n')
    paragraphs = text.strip().split('\n\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("[Your Name]")
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_text_docx(text: str, title: str = "Document", font_family: str = "Times New Roman", font_size: int = 11, template: str = "standard") -> io.BytesIO:
    """Create a Word document from raw text, preserving basic formatting and styling"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_family
    font.size = Pt(font_size)

    is_modern = template == 'modern'
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Split text into lines and add to document
    # PRE-PROCESS: Replace literal \n or \\n strings from LLM with real newlines
    text = (text or "").replace('\\n', '\n').replace('\\\\n', '\n')
    lines = text.strip().split('\n')
    
    # First line is usually the name, second might be the title
    if lines:
        name_text = lines[0].strip()
        if name_text.upper().startswith("NAME:"):
            name_text = name_text[5:].strip()
            
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name_text.upper())
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.name = "Arial"
        if not is_modern:
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(2)
        
        lines = lines[1:]
        
        # Check if second line is a job title (not a section header or bullet)
        if lines and lines[0].strip() and not lines[0].strip().startswith('=') and not lines[0].strip().startswith('•'):
            title_text = sanitize_job_title(lines[0].strip())
            if title_text:
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(title_text.upper())
                title_run.bold = False
                title_run.font.size = Pt(11)
                title_run.font.name = "Arial"
                title_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                if not is_modern:
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_para.paragraph_format.space_after = Pt(12)
                lines = lines[1:]

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
            
        if line_stripped.startswith('===') and line_stripped.endswith('==='):
            # Section header
            p = doc.add_paragraph()
            clean_title = line_stripped.replace('=', '').strip().upper()
            
            # Skip literal generic headers that AI might provide
            if clean_title in ["NAME", "CONTACT"]:
                continue
                
            run = p.add_run(clean_title)
            run.bold = True
            run.font.size = Pt(11)
            if not is_modern:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_bottom_border(p)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif (line_stripped.isupper() and len(line_stripped) < 50) or (line_stripped.startswith('#') and len(line_stripped) < 60) or (line_stripped.upper() in ["EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EDUCATION", "PROJECTS", "SKILLS", "SUMMARY", "PROFESSIONAL SUMMARY", "PROFILE", "CERTIFICATIONS", "AWARDS", "LANGUAGES", "TECHNICAL SKILLS", "CORE COMPETENCIES"]):
            # Main sections like PROFESSIONAL EXPERIENCE
            clean_title = line_stripped.replace('#', '').strip().upper()
            
            # Skip literal generic headers
            if clean_title in ["NAME", "CONTACT"]:
                continue
                
            p = doc.add_paragraph()
            run = p.add_run(clean_title)
            run.bold = True
            run.font.size = Pt(11)
            if not is_modern:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_bottom_border(p)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        elif line_stripped.startswith('- ') or line_stripped.startswith('• '):
            # Bullet point
            p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
        else:
            # Smarter header detection logic (Sync with Frontend ResumePaper.jsx)
            has_separator = (" — " in line_stripped or " | " in line_stripped or " – " in line_stripped)
            # Date detection: matches years like 2024 or months like Jan/February
            import re
            has_date = bool(re.search(r'\d{4}', line_stripped)) or bool(re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', line_stripped, re.I))
            
            # Line is a header if it's potentially an entry start (prev line empty), 
            # or if it has info like job title/company/dates
            # We also ensure it doesn't end with a period (unlikely for a header)
            is_header = (has_separator and len(line_stripped) < 150) or (has_date and len(line_stripped) < 100)
            
            p = doc.add_paragraph()
            if is_header and not line_stripped.endswith('.'):
                run = p.add_run(line)
                run.bold = True
            else:
                # Try to center contact info if it's the first non-empty line after the name
                # Strip literal "CONTACT:" prefix if AI included it
                if i == 0 and line_stripped.upper().startswith("CONTACT:"):
                     p.clear()
                     p.add_run(line_stripped[8:].strip())
                     
                if i == 0 and ('@' in line or '|' in line) and not is_modern:
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.add_run(line)

            p.paragraph_format.space_after = Pt(0)
            
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

async def refine_resume_section(section_name: str, section_content: str, job_description: str, resume_context: str = "") -> dict:
    """
    Selectively improves a specific section of the resume using AI.
    """
    logger.info(f"Refining section: {section_name}")
    
    prompt = f"""You are an Expert Resume Writer. Optimize the following resume section specifically for the target job description.

TARGET JOB DESCRIPTION:
{job_description[:2000]}

SECTION TO REFINE: {section_name.upper()}
CURRENT CONTENT:
{section_content}

FULL RESUME CONTEXT (for consistency):
{resume_context[:2000]}

MARS TAILORING RULES:
1. STRIP MARKDOWN: Output only clean plain text. Zero markdown (**bold**, *italics*, etc).
2. TIMELINE LAW: Only assign technology/tools that were in mainstream use DURING the dates for this content.
3. BANNED WORDS: Do NOT use "passionate", "motivated", "excited", "leverage", "synergy", "vibe coding".
4. SENIOR FRAMING: For experience/projects, use judgment language ("designed for failure tolerance", "reasoned about edge cases").
5. PRESERVE FACTS: Do not invent metrics, roles, or degrees.
6. FORMAT: Respond with the refined text ONLY. No prose. No markdown markers.

REFINED {section_name.upper()} CONTENT:"""

    try:
        refined_text = await unified_api_call(
            prompt,
            max_tokens=1000,
            model="llama-3.3-70b-versatile"
        )
        
        if not refined_text:
            return {"error": "AI failed to generate refined content"}

        # Cleanup
        refined_text = refined_text.strip()
        # Remove markdown markers that might have leaked
        refined_text = re.sub(r'[*_#>]', '', refined_text)
        
        return {"refined_content": refined_text}
    except Exception as e:
        logger.error(f"Error in refine_resume_section: {e}")
        return {"error": str(e)}
