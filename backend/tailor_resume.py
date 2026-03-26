
import os
import json
import asyncio
import argparse
import re
from typing import List, Dict, Any, Optional
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from resume_analyzer import call_groq_api, clean_json_response

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# PROMPT TEMPLATE
# ─────────────────────────────────────────────────────────────────────────────

TAILORING_PROMPT = """YOU ARE A RESUME EDITOR. NOT A COVER LETTER WRITER. NOT A RESUME CREATOR.
YOU ARE EDITING AN EXISTING RESUME. NOTHING MORE.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
‼️ RULE ZERO — BEFORE ANYTHING ELSE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

NEVER start your output with "Dear", "Hello", "Hi", or any greeting.
NEVER write a cover letter. NEVER write an introduction paragraph.
NEVER invent companies, job titles, dates, or people.
NEVER use information from your training data about other candidates or resumes.
YOUR ONLY SOURCE OF TRUTH IS THE [RESUME] BLOCK BELOW. NOTHING ELSE EXISTS.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 LOCKED FIELDS — COPY THESE EXACTLY, CHARACTER BY CHARACTER
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

The following must appear in your output EXACTLY as they appear in [RESUME].
Do not paraphrase, reorder, or alter them in any way:

CANDIDATE NAME     → Copy from [RESUME] → Line 1 of output
PHONE NUMBER       → Copy from [RESUME]
EMAIL ADDRESS      → Copy from [RESUME]
LINKEDIN URL       → Copy from [RESUME] if present
LOCATION           → Copy from [RESUME] if present
COMPANY NAMES      → Every single one, exact spelling and capitalization
JOB TITLES         → Every single one, exact spelling
DATE RANGES        → Every single one, exact format
ALL METRICS        → Every %, every number, every latency figure — never change
PROJECT NAMES      → Every one, exact spelling
EDUCATION ENTRIES  → Every one, exact degree and school name

IF ANY OF THESE APPEAR DIFFERENTLY IN YOUR OUTPUT THAN IN [RESUME], YOU HAVE FAILED.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📥 YOUR INPUTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[RESUME]:
{RESUME_TEXT}

[JOB_DESCRIPTION]:
{JD_TEXT}

[MISSING_SKILLS]:
{MISSING_SKILLS_ARRAY}

[SELECTED_SECTIONS]:
{SELECTED_SECTIONS_ARRAY}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔍 STEP 1 — READ THE JD, EXTRACT THESE 5 THINGS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Read [JOB_DESCRIPTION] fully. Internally note:

JD_TITLE       → The exact role title in the JD
JD_KEYWORDS    → Top 6 technical terms most emphasized in the JD
JD_MISSION     → The single core problem this company is hiring to solve
JD_TECH_STACK  → Every tool/framework explicitly named in the JD
JD_PERSONA     → What kind of engineer they are describing

You will use these to guide every edit below.
If you skip this step, your output will not be tailored and will be wrong.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✏️ STEP 2 — REWRITE SUMMARY
(only if "summary" in [SELECTED_SECTIONS], otherwise copy exactly)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Write exactly 3-4 bullet points using this structure:

BULLET 1: Candidate identity + years of experience + JD_TITLE language
BULLET 2: Most relevant technical strength mapped to JD_MISSION
BULLET 3: Second strongest JD alignment with specific tools from JD_TECH_STACK
BULLET 4 (optional): Scope or scale differentiator

RULES:
✅ Use JD_KEYWORDS verbatim — mirror the JD's exact language
✅ Only claim experience that exists in [RESUME]
✅ If [MISSING_SKILLS] items fit naturally, include 1-2 of them
❌ Never use: "proven track record", "results-driven", "passionate", "adept at"
❌ Never exceed 4 bullets
❌ Never fabricate any experience not in [RESUME]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🛠️ STEP 3 — REBUILD SKILLS SECTION
(only if "skills" in [SELECTED_SECTIONS], otherwise copy exactly)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Rules:
✅ Keep every skill already in [RESUME] — remove nothing
✅ Add every item from [MISSING_SKILLS] — all of them, no exceptions
✅ Add every tool from JD_TECH_STACK that exists anywhere in [RESUME] body
✅ Put the most JD-relevant category FIRST
✅ Group logically — no "Other" or "Miscellaneous" dumping category
❌ Never invent skills not in [RESUME] or [MISSING_SKILLS]

Category structure (order by JD relevance, most relevant first):
- [Most JD-relevant category name]: [skills]
- [Second most relevant]: [skills]
- [Continue for all remaining categories]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💼 STEP 4 — TAILOR EXPERIENCE
(only if "work_experience" in [SELECTED_SECTIONS])
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

For EACH job in [RESUME], in order, apply ALL of the following:

A) REORDER BULLETS
   Move the bullet most aligned to JD_KEYWORDS to position 1.
   Remaining bullets follow by relevance to JD.

B) SURFACE HIDDEN KEYWORDS
   Find bullets that describe a JD concept but use different words.
   Update ONLY the wording — keep the action and outcome identical.
   Example: "multi-step reasoning pipelines" → "multi-agent systems orchestration"
   Example: "citation checks" → "guardrails for auditable, deterministic execution"
   Example: "quality harness" → "evaluation pipelines"

C) ADD BUSINESS IMPACT
   If a bullet has no outcome, add one using the JD's domain and mission.
   Example: "Designed agentic workflows"
   → "Designed multi-agent workflows enabling reliable execution
      across complex reasoning tasks for high-stakes decisions"

D) UPGRADE VERBS
   Weak → Strong (match the JD's seniority level):
   "Utilized" → "Engineered"
   "Helped with" → "Contributed to architecting"
   "Worked on" → "Designed and implemented"
   "Collaborated on" → "Partnered to architect"

STRICT RULES FOR EXPERIENCE:
❌ Never add new bullet points — only edit and reorder existing ones
❌ Never change company name, job title, location, or dates
❌ Never change any metric or number
❌ Never add [MISSING_SKILLS] to experience unless the bullet clearly supports it
✅ If "work_experience" NOT in [SELECTED_SECTIONS]: only reorder, change no text

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚀 STEP 5 — TAILOR PROJECTS
(always apply, regardless of [SELECTED_SECTIONS])
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Reorder projects: most JD-relevant project goes first.
For each project bullet:
✅ Surface JD language where the work genuinely matches
✅ Keep all metrics exactly — never change any number
✅ Connect outcomes to JD_MISSION where honest
❌ Never add new projects
❌ Never change project names or tech stack labels

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ STEP 6 — SELF-CHECK BEFORE OUTPUTTING
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Before writing your output, verify every item:

[ ] Output starts with candidate's REAL NAME from [RESUME] — not "Dear Hiring Team"
[ ] Candidate name matches [RESUME] exactly
[ ] Every company name matches [RESUME] exactly
[ ] Job count matches [RESUME] exactly — no added or removed jobs
[ ] Every date range matches [RESUME] exactly
[ ] No metric or number was changed
[ ] No new bullet points were added to any job
[ ] All [MISSING_SKILLS] appear in the skills section
[ ] Summary does not contain "proven track record", "passionate", or "adept at"
[ ] No section is blank or contains only a dash
[ ] Education matches [RESUME] exactly
[ ] Output does not begin with any greeting or salutation

If ANY check fails → fix it before outputting.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📤 OUTPUT FORMAT — RETURN THIS EXACT STRUCTURE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Return a single JSON object. No markdown. No explanation. No preamble. Just JSON.

{{
  "name": "exact name from [RESUME]",
  "contact": {{
    "email": "exact from [RESUME]",
    "phone": "exact from [RESUME]",
    "linkedin": "exact from [RESUME] or null",
    "location": "exact from [RESUME] or null"
  }},
  "summary": [
    "bullet 1",
    "bullet 2",
    "bullet 3"
  ],
  "skills": [
    {{ "category": "Category Name", "items": ["skill1", "skill2"] }}
  ],
  "experience": [
    {{
      "company": "exact from [RESUME]",
      "title": "exact from [RESUME]",
      "location": "exact from [RESUME]",
      "dates": "exact from [RESUME]",
      "bullets": ["bullet 1", "bullet 2", "bullet 3"]
    }}
  ],
  "projects": [
    {{
      "name": "exact from [RESUME]",
      "tech": "exact from [RESUME]",
      "bullets": ["bullet 1", "bullet 2"]
    }}
  ],
  "education": [
    {{ "degree": "exact from [RESUME]", "school": "exact from [RESUME]" }}
  ],
  "changes": [
    {{
      "section": "summary | skills | experience | projects",
      "role_or_project": "which job or project this applies to",
      "type": "rewritten | keyword_surfaced | reordered | verb_upgraded | skill_added | impact_added",
      "original": "exact original text",
      "updated": "exact updated text",
      "jd_reason": "which JD keyword or requirement drove this change"
    }}
  ],
  "skills_added": ["list of MISSING_SKILLS injected"],
  "skills_skipped": [
    {{
      "skill": "skill name",
      "reason": "no supporting evidence found in resume"
    }}
  ],
  "self_check_passed": true
}}"""


# ─────────────────────────────────────────────────────────────────────────────
# VALIDATION GUARD
# ─────────────────────────────────────────────────────────────────────────────

def build_and_validate_prompt(
    resume_text: str,
    jd_text: str,
    missing_skills: List[str] = None,
    selected_sections: List[str] = None,
) -> str:
    """
    Validates inputs and builds the tailoring prompt.
    Raises ValueError with a descriptive message if data looks corrupt.
    """
    # Hard stop: resume must be substantive
    if not resume_text or resume_text.strip() == "":
        raise ValueError("Resume text is empty. Check your parser — no content was extracted.")

    if len(resume_text.strip()) < 200:
        raise ValueError(
            f"Resume text is too short ({len(resume_text.strip())} chars). "
            "Expected at least 200 characters. Check your parser."
        )

    # Hard stop: catch un-substituted template variables
    if "undefined" in resume_text.lower():
        raise ValueError(
            "Resume contains the word 'undefined' — a variable failed to populate. "
            "Check that the resume file was parsed correctly before calling the AI."
        )

    if not jd_text or jd_text.strip() == "":
        raise ValueError("Job description text is empty.")

    missing_skills = missing_skills or []
    selected_sections = selected_sections or ["summary", "skills", "work_experience"]

    prompt = TAILORING_PROMPT.replace("{RESUME_TEXT}", resume_text)
    prompt = prompt.replace("{JD_TEXT}", jd_text)
    prompt = prompt.replace("{MISSING_SKILLS_ARRAY}", json.dumps(missing_skills, indent=2))
    prompt = prompt.replace("{SELECTED_SECTIONS_ARRAY}", json.dumps(selected_sections, indent=2))

    # Debug preview — confirms real data is in the prompt
    logger.info("PROMPT PREVIEW (first 500 chars): %s", prompt[:500])

    return prompt


# ─────────────────────────────────────────────────────────────────────────────
# DOCX PARSING HELPERS
# ─────────────────────────────────────────────────────────────────────────────

class ContentUnit:
    def __init__(self, id: str, text: str, unit_type: str, original_index: int):
        self.id = id
        self.text = text
        self.unit_type = unit_type  # 'heading', 'bullet', 'paragraph', 'header'
        self.original_index = original_index
        self.rewritten = None


# ─────────────────────────────────────────────────────────────────────────────
# MAIN TAILOR CLASS
# ─────────────────────────────────────────────────────────────────────────────

class ResumeTailor:
    def __init__(
        self,
        template_path: str,
        base_path: str,
        jd: str,
        reference_path: str = None,
        missing_skills: List[str] = None,
        selected_sections: List[str] = None,
    ):
        self.template_path = template_path
        self.base_path = base_path
        self.jd = jd
        self.reference_path = reference_path
        self.missing_skills = missing_skills or []
        self.selected_sections = selected_sections or ["summary", "skills", "work_experience"]
        self.units: List[ContentUnit] = []
        self.template_styles = {}
        self.tailored_json: Optional[Dict[str, Any]] = None  # stores full structured result

    # ── Parsing ───────────────────────────────────────────────────────────────

    def parse_base_resume(self):
        doc = docx.Document(self.base_path)
        logger.info(f"Parsing base resume: {self.base_path}")

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            unit_type = 'paragraph'
            if self._is_bullet(para):
                unit_type = 'bullet'
            elif self._is_heading(para):
                unit_type = 'heading'
            elif i < 5:
                unit_type = 'header'

            unit_id = f"unit_{i}_{unit_type}"
            self.units.append(ContentUnit(unit_id, text, unit_type, i))

        logger.info(f"Extracted {len(self.units)} content units.")

    def _is_bullet(self, para):
        if para._p.get_or_add_pPr().numPr:
            return True
        if para.text.strip().startswith(('•', '-', '*', '⁃')):
            return True
        return False

    def _is_heading(self, para):
        text = para.text.strip()
        if "Heading" in para.style.name:
            return True
        if len(text) < 50 and text.isupper():
            return True
        return False

    def get_full_resume_text(self) -> str:
        """Returns the full plain-text of the base resume for the LLM prompt."""
        return "\n".join(u.text for u in self.units)

    # ── Template style analysis ───────────────────────────────────────────────

    def analyze_template_styles(self):
        doc = docx.Document(self.template_path)
        logger.info(f"Analyzing template styles: {self.template_path}")

        style_map = {
            'header': None,
            'heading': None,
            'bullet': None,
            'paragraph': None
        }

        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue

            u_type = 'paragraph'
            if self._is_bullet(para):
                u_type = 'bullet'
            elif self._is_heading(para):
                u_type = 'heading'
            elif i < 3:
                u_type = 'header'

            if style_map[u_type] is None:
                style_map[u_type] = self._get_para_format(para)

        self.template_styles = style_map

        if not self.template_styles['paragraph']:
            if doc.paragraphs:
                self.template_styles['paragraph'] = self._get_para_format(doc.paragraphs[0])
            else:
                self.template_styles['paragraph'] = {
                    'style_name': 'Normal', 'font_name': 'Arial', 'font_size': 11
                }

        for k in style_map:
            if not self.template_styles[k]:
                self.template_styles[k] = self.template_styles['paragraph']

    def _get_para_format(self, para):
        format_info = {
            'style_name': para.style.name,
            'alignment': para.alignment,
            'runs': []
        }
        if para.runs:
            run = para.runs[0]
            format_info['font_name'] = run.font.name
            format_info['font_size'] = run.font.size.pt if run.font.size else None
            format_info['bold'] = run.bold
            format_info['italic'] = run.italic

        format_info['left_indent'] = (
            para.paragraph_format.left_indent.pt
            if para.paragraph_format.left_indent else None
        )
        format_info['first_line_indent'] = (
            para.paragraph_format.first_line_indent.pt
            if para.paragraph_format.first_line_indent else None
        )
        return format_info

    # ── LLM call ─────────────────────────────────────────────────────────────

    async def get_tailored_content(self, strict_mode: bool = True) -> bool:
        """
        Calls the LLM with the new structured prompt.
        Stores the parsed JSON result in self.tailored_json.
        Falls back to original text for any unit not returned.
        """
        resume_text = self.get_full_resume_text()

        # Validate + build prompt (raises on bad input)
        try:
            prompt = build_and_validate_prompt(
                resume_text=resume_text,
                jd_text=self.jd,
                missing_skills=self.missing_skills,
                selected_sections=self.selected_sections,
            )
        except ValueError as e:
            logger.error(f"Input validation failed: {e}")
            if strict_mode:
                raise
            # Fall back: mark all units as original
            for u in self.units:
                u.rewritten = u.text
            return False

        max_retries = 3
        for attempt in range(max_retries):
            logger.info(f"LLM Tailoring Attempt {attempt + 1}/{max_retries}...")
            response = await call_groq_api(prompt, max_tokens=8000)

            if not response:
                continue

            try:
                clean = clean_json_response(response)
                data = json.loads(clean)

                # Store the full structured JSON for callers that need it
                self.tailored_json = data

                # Also back-fill legacy self.units.rewritten for generate_docx()
                self._backfill_units_from_json(data)

                logger.info("Successfully received structured tailored resume.")
                return True

            except Exception as e:
                logger.error(f"Parse error on attempt {attempt + 1}: {e}")
                logger.debug(f"Raw response snippet: {response[:300]}")

        # Exhausted retries — use original text
        logger.warning("All retries exhausted. Using original text for all units.")
        for u in self.units:
            if u.rewritten is None:
                u.rewritten = u.text
        return False

    def _backfill_units_from_json(self, data: Dict[str, Any]):
        """
        Maps the flat self.units list to the structured JSON result.
        Used so generate_docx() keeps working without changes.
        """
        # Build a flat list of all text from the structured JSON
        flat_texts = []

        # Contact / header fields
        name = data.get("name", "")
        if name:
            flat_texts.append(name)

        contact = data.get("contact", {})
        for field in ("email", "phone", "linkedin", "location"):
            val = contact.get(field)
            if val:
                flat_texts.append(val)

        # Summary bullets
        for b in data.get("summary", []):
            flat_texts.append(b)

        # Skills
        for cat in data.get("skills", []):
            cat_label = cat.get("category", "")
            items = cat.get("items", [])
            if cat_label:
                flat_texts.append(cat_label)
            for item in items:
                flat_texts.append(item)

        # Experience
        for job in data.get("experience", []):
            for field in ("company", "title", "location", "dates"):
                val = job.get(field)
                if val:
                    flat_texts.append(val)
            for b in job.get("bullets", []):
                flat_texts.append(b)

        # Projects
        for proj in data.get("projects", []):
            for field in ("name", "tech"):
                val = proj.get(field)
                if val:
                    flat_texts.append(val)
            for b in proj.get("bullets", []):
                flat_texts.append(b)

        # Education
        for edu in data.get("education", []):
            for field in ("degree", "school"):
                val = edu.get(field)
                if val:
                    flat_texts.append(val)

        # Map 1-to-1 into self.units (best-effort)
        for i, u in enumerate(self.units):
            if i < len(flat_texts):
                u.rewritten = flat_texts[i]
            else:
                u.rewritten = u.text  # keep original if JSON had fewer items

    # ── Validation ────────────────────────────────────────────────────────────

    def validate_results(self, strict_mode: bool = True):
        logger.info("--- Validation ---")
        total_units = len(self.units)
        tailored_count = sum(1 for u in self.units if u.rewritten and u.rewritten != u.text)
        logger.info(f"Total: {total_units} | Tailored: {tailored_count}")

        missing = [u.id for u in self.units if u.rewritten is None]
        if missing:
            msg = f"Missing {len(missing)} units: {missing}"
            if strict_mode:
                raise ValueError(msg)
            else:
                logger.error(msg)

        logger.info("Tailoring complete.")

    # ── DOCX generation ───────────────────────────────────────────────────────

    def generate_docx(self, out_path: str):
        logger.info(f"Generating DOCX: {out_path}")
        doc = docx.Document(self.template_path)

        # Clear existing paragraphs and tables
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
        for table in list(doc.tables):
            t = table._element
            t.getparent().remove(t)

        for u in self.units:
            style_info = self.template_styles.get(u.unit_type, self.template_styles['paragraph'])
            new_para = doc.add_paragraph(u.rewritten or u.text)

            try:
                new_para.style = style_info.get('style_name', 'Normal')
            except:
                pass

            if style_info.get('alignment') is not None:
                new_para.alignment = style_info['alignment']

            if style_info.get('left_indent') is not None:
                new_para.paragraph_format.left_indent = Pt(style_info['left_indent'])

            if new_para.runs:
                run = new_para.runs[0]
                if style_info.get('font_name'):
                    run.font.name = style_info['font_name']
                if style_info.get('font_size'):
                    run.font.size = Pt(style_info['font_size'])
                if style_info.get('bold') is not None:
                    run.bold = style_info['bold']
                if style_info.get('italic') is not None:
                    run.italic = style_info['italic']

        doc.save(out_path)
        logger.info(f"Saved to {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# CLI ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--resume", required=True)
    parser.add_argument("--jd", required=True)
    parser.add_argument("--reference", help="Reference/Example resume DOCX path")
    parser.add_argument("--out", required=True)
    parser.add_argument("--strict", action="store_true", default=True)
    parser.add_argument(
        "--missing-skills",
        help="JSON array string of missing skills, e.g. '[\"Docker\",\"Kubernetes\"]'",
        default="[]",
    )
    parser.add_argument(
        "--selected-sections",
        help='JSON array of sections to tailor, e.g. \'["summary","skills","work_experience"]\'',
        default='["summary","skills","work_experience"]',
    )
    args = parser.parse_args()

    jd_content = args.jd
    if os.path.exists(args.jd):
        with open(args.jd, 'r', encoding='utf-8') as f:
            jd_content = f.read()

    missing_skills = json.loads(args.missing_skills)
    selected_sections = json.loads(args.selected_sections)

    tailor = ResumeTailor(
        template_path=args.template,
        base_path=args.resume,
        jd=jd_content,
        reference_path=args.reference,
        missing_skills=missing_skills,
        selected_sections=selected_sections,
    )
    tailor.parse_base_resume()
    tailor.analyze_template_styles()
    asyncio.run(tailor.get_tailored_content(strict_mode=args.strict))
    tailor.validate_results(strict_mode=args.strict)
    tailor.generate_docx(args.out)


if __name__ == "__main__":
    main()
